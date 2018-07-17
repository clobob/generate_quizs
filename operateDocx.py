#!--encoding=utf-8--!

'''
instruction: this script is used to write word document.
dependent: py-docx 

'''
from docx import Document
from docx.shared import Inches


'''
reference method, not in use.
'''
def operDocx():
    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='IntenseQuote')

    document.add_paragraph(
        'first item in unordered list', style='ListBullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='ListNumber'
    )

#    document.add_picture('monty-truth.png', width=Inches(1.25))

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    # for item in recordset:
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = str(item.qty)
    #     row_cells[1].text = str(item.id)
    #     row_cells[2].text = item.desc

    document.add_page_break()

    document.save('demo.docx')

def writeDocx(contentList,docxName):
    document = Document()
    # p = document.add_paragraph(content)
    # document.add_page_break()
    total = contentList[0].__len__()
    print 'amount of content: '+ str(total)
    p = 0
    q = 0
    hasContent = True
    # row = 20
    # col = 3
    t = 1
    while(hasContent):
        if p >= total:
            hasContent = False
            break
        # if total - p > 60:

        table = document.add_table(rows=20, cols=3)
        # else:
        #     row = (total -p)//3 + 1
        #     table = document.add_table(rows=20, cols=3)
        print 'table added  ' + str(t)
        num = 1
        for i in range(20):
            # print 'insert row data'
            for j in range(3):
                if p >= total:
                    hasContent = False
                    break
                table.rows[i].cells[j].text = str(num) + ')\t' + contentList[0][p]
                num = num +1
                p = p + 1
        # document.add_paragraph('----------------------------------------------------------------------------------------------------------------------')
        # document.add_paragraph('----------------------------------------------------------------------------------------------------------------------')
        # table = document.add_table(rows=3, cols=20)
        # for m in range(3):
        #     for n in range(20):
        #         table.rows[m].cells[n].text = str(q + 1) +')'+ str(contentList[1][q])
        #         q = q + 1

#        results = ''
#        num = 1
#        for m in range(60):
#            results = results +'^'+ str(num)+'>'+str(contentList[1][q])+'$'
#            q = q + 1
#            num +=1
#        par = document.add_paragraph(u'本页答案： ')
#        par.add_run(results).italic = True
        print 'now in the '+ str(p) + 'step'
        t = t +1
        # document.add_paragraph('----------------------------------------------------------------------------------------------------------------------')
        document.add_page_break()
    document.save(docxName)

# def id(id=0):
#     return id+1

def saveExpression(Expressions,docxName):
    document = Document()
    total = len(Expressions)
    print 'Total Expressions are: '+str(total)
    
    index=0
    while index<total:
        table = document.add_table(rows=25, cols=4)
        for i in range(25):
            for j in range(4):
                table.rows[i].cells[j].text = Expressions[index]
                index = index+1
        document.add_page_break()
    document.save(docxName)

if __name__=="__main__":
    print 'this is the main method'
    operDocx()
    # order = id()
    # writeDocx( '\t'+str(order) + '.'+ ' 12  +  44 =\t' + str(order) + '.'+ '\t 12  +  44 =\t' + str(order) + '.'+ '\t 12  +  44 =\t'+'\n\r','quiz.docx')
